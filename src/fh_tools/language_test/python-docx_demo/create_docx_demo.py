#! /usr/bin/env python
# -*- coding:utf-8 -*-
"""
@author  : MG
@Time    : 19-5-10 下午1:37
@File    : create_docx_demo.py
@contact : mmmaaaggg@163.com
@desc    : https://python-docx.readthedocs.io/en/latest/
"""
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


def main():
    # 创建文档对象
    document = Document()

    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 创建自定义字符样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle2 = document.styles.add_style('UserStyle2', 2)
    # 设置字体尺寸
    UserStyle2.font.size = Pt(15)
    # 设置字体颜色0c8ac5
    UserStyle2.font.color.rgb = RGBColor(0x0c, 0x8a, 0xc5)
    # 设置段落样式为宋体
    UserStyle2.font.name = '宋体'
    UserStyle2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 使用自定义段落样式
    document.add_paragraph('自定义段落样式', style=UserStyle1)

    # 使用自定义字符样式
    document.add_paragraph('').add_run('正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。', style=UserStyle2)

    # 设置粗体字
    document.add_paragraph('设置粗体字:').add_run('粗体字').bold = True

    # 设置斜体字
    document.add_paragraph('设置斜体字:').add_run('斜体字').italic = True

    # 设置字号50
    document.add_paragraph('设置字号50:').add_run('50').font.size = Pt(50)

    # 设置字体颜色为 af2626
    document.add_paragraph('设置字体颜色:').add_run('颜色').font.color.rgb = RGBColor(0xaf, 0x26, 0x26)

    # 样式叠加: 将字体改到30号并且将字体改成特定颜色;
    doubleStyle = document.add_paragraph('同时设置文字颜色和字号:').add_run('颜色和尺寸')
    doubleStyle.font.size = Pt(30)
    doubleStyle.font.color.rgb = RGBColor(0xaf, 0x26, 0x26)

    # 添加分页符
    document.add_page_break()

    # 创建 有序列表
    document.add_paragraph('').add_run('有序列表').font.size = Pt(30)
    document.add_paragraph('把冰箱门打开', style='List Number')
    document.add_paragraph('把大象装进去', style='List Number')
    document.add_paragraph('把冰箱门关上', style='List Number')

    # 创建 无序列表
    document.add_paragraph('').add_run('无序列表').font.size = Pt(30)
    document.add_paragraph('天地匆匆 惊鸿而过 路有千百个', style='List Bullet')
    document.add_paragraph('遑遑无归 闲云逸鹤 人间红尘过', style='List Bullet')
    document.add_paragraph('引势而流 鸿门乱局 各有各选择', style='List Bullet')
    document.add_paragraph('乾震坎艮 坤巽离兑 定一切生克', style='List Bullet')

    # 添加分页符
    document.add_page_break()
    # 添加图片
    document.add_paragraph('').add_run('添加图片').font.size = Pt(30)
    document.add_picture('少女17087938.jpg', width=Inches(5))

    # 添加分页符
    document.add_page_break()

    document.add_paragraph('').add_run('创建表格').font.size = Pt(30)
    # 创建两行两列的表格
    rows_num = 5
    cols_num = 6
    table = document.add_table(rows=rows_num, cols=cols_num, style='Table Grid')

    for r in range(rows_num):
        for c in range(cols_num):
            table.cell(r, c).text = "第{r}行{c}列".format(r=r + 1, c=c + 1)

    document.add_paragraph('pd.DataFrame 生成表格')
    df = pd.DataFrame({'a': [1, 2, 3], 'b': [2, 3, 4], 'c': [3, 4, 5], 'd': [4, 5, 6]})
    df_2_table(document, df)

    # 保存文档
    document.save('create_docx_demo.docx')


def df_2_table(doc, df):
    row_num, col_num = df.shape
    t = doc.add_table(row_num + 1, col_num)
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for j in range(col_num):
        # t.cell(0, j).text = df.columns[j]
        # paragraph = t.cell(0, j).add_paragraph()
        paragraph = t.cell(0, j).paragraphs[0]
        paragraph.add_run(df.columns[j]).bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # format table style to be a grid
    t.style = 'TableGrid'

    # populate the table with the dataframe
    for i in range(row_num):
        for j in range(col_num):
            # t.cell(i + 1, j).text = str(df.values[i, j])
            paragraph = t.cell(i + 1, j).paragraphs[0]
            paragraph.add_run(str(df.values[i, j])).bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Highlight all cells limegreen (RGB 32CD32) if cell contains text "0.5"
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    for i in range(1, row_num + 1):
        for j in range(col_num):
            if i % 2 == 0:
                t.cell(i, j)._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="A3D9EA"/>'.format(nsdecls('w'))))


if __name__ == '__main__':
    main()
